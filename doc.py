import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# 训练数据
data = {
    "num_runs": 1,
    "total_rounds": 10,
    "runs": [
        {
            "cnn_acc": [0.711, 0.889, 0.908, 0.913, 0.91, 0.93, 0.938, 0.915, 0.914, 0.913],
            "tf_acc": [0.72, 0.88, 0.801, 0.889, 0.887, 0.894, 0.885, 0.871, 0.879, 0.882],
            "ensemble_acc": [0.71, 0.888, 0.894, 0.9, 0.898, 0.903, 0.904, 0.886, 0.89, 0.891]
        }
    ],
    "config": {
        "CV_TH": 0.02,
        "ARI_TH": 0.24,
        "BATCH_SIZE": 32,
        "NUM_EPOCHS": 60,
        "LR": 0.001,
        "TH_HIGH": 0.95,
        "TH_LOW": 0.05
    }
}

# 生成训练准确率曲线图
def generate_training_plot():
    rounds = list(range(1, 11))
    cnn_acc = data['runs'][0]['cnn_acc']
    tf_acc = data['runs'][0]['tf_acc']
    ensemble_acc = data['runs'][0]['ensemble_acc']
    
    plt.figure(figsize=(10, 6))
    plt.plot(rounds, cnn_acc, 'b-o', label='CNN分支准确率', linewidth=2, markersize=6)
    plt.plot(rounds, tf_acc, 'r-s', label='Transformer分支准确率', linewidth=2, markersize=6)
    plt.plot(rounds, ensemble_acc, 'g-^', label='集成模型准确率', linewidth=2, markersize=6)
    
    plt.xlabel('训练轮次 (Round)', fontsize=12)
    plt.ylabel('准确率 (Accuracy)', fontsize=12)
    plt.title('ConvTransformer架构训练准确率曲线', fontsize=14, fontweight='bold')
    plt.legend(fontsize=10)
    plt.grid(True, alpha=0.3)
    plt.xticks(rounds)
    plt.ylim(0.6, 1.0)
    
    # 在数据点上标注数值
    for i, (cnn, tf, ens) in enumerate(zip(cnn_acc, tf_acc, ensemble_acc)):
        plt.annotate(f'{cnn:.3f}', (rounds[i], cnn), textcoords="offset points", xytext=(0,5), ha='center', fontsize=8, color='blue')
        plt.annotate(f'{tf:.3f}', (rounds[i], tf), textcoords="offset points", xytext=(0,5), ha='center', fontsize=8, color='red')
        plt.annotate(f'{ens:.3f}', (rounds[i], ens), textcoords="offset points", xytext=(0,5), ha='center', fontsize=8, color='green')
    
    plt.tight_layout()
    plt.savefig('training_accuracy_curve.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("图片已生成：training_accuracy_curve.png")

# 生成Word文档
def generate_word_doc():
    doc = Document()
    
    # 标题
    title = doc.add_heading('周报：ConvTransformer架构优化与实验进展', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('汇报人：').bold = True
    p.add_run('[您的姓名]\n')
    p.add_run('汇报周期：').bold = True
    p.add_run('2026年3月第2周\n')
    p.add_run('项目：').bold = True
    p.add_run('ECG房颤检测模型研发')
    
    # 一、本周工作概述
    doc.add_heading('一、本周工作概述', level=1)
    doc.add_paragraph('本周重点完成了 ConvTransformer 架构的深度优化，针对原始纯Transformer在ECG长序列建模中的局限性，引入"局部卷积+全局注意力"的混合机制，并完成了首轮训练验证。主要工作包括：')
    
    items = [
        '✅ 重构模型主干：实现 1D-CNN + Transformer Encoder 的级联架构',
        '✅ 引入 [CLS] Token 机制替代全局平均池化，增强稀疏病理特征捕捉能力',
        '✅ 优化降采样策略：卷积核调整为 7→5→3 渐进式，保留时序分辨率',
        '✅ 升级位置编码：由固定三角函数改为可学习参数矩阵（Learnable Positional Encoding）',
        '✅ 完成首轮训练（10 rounds），采集CNN/Transformer/Ensemble三路准确率指标'
    ]
    
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 二、架构优化核心要点
    doc.add_heading('二、架构优化核心要点', level=1)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '优化项'
    hdr_cells[1].text = '原方案问题'
    hdr_cells[2].text = '改进方案'
    hdr_cells[3].text = '预期收益'
    
    # 加粗表头
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 表格数据
    table_data = [
        ('特征提取', '纯Transformer缺乏局部归纳偏置，难捕捉P/QRS/T波微观形态', '前置多层1D-CNN+MaxPool，充当"信号放大镜"', '提升局部波形特征表达能力'),
        ('序列聚合', '全局平均池化易稀释稀疏的房颤异常特征', '引入可学习 [CLS] Token，通过注意力主动吸附关键节点', '避免特征平均化，增强病理信号敏感度'),
        ('时序建模', '固定Sin/Cos编码对规整高采样ECG信号适配性弱', '改为 nn.Parameter 可学习位置编码，反向传播自适应微调', '提升心跳间相对/绝对依赖关系的建模精度'),
        ('降采样策略', '激进池化导致时序信息过度压缩，Attention计算空间不足', '缩小卷积核、取消大步长，仅靠MaxPool稳态降维', '为Transformer保留充足时序分辨率（数百步）')
    ]
    
    for item, problem, solution, benefit in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = problem
        row_cells[2].text = solution
        row_cells[3].text = benefit
    
    # 三、实验结果分析
    doc.add_heading('三、实验结果分析', level=1)
    
    doc.add_heading('📊 训练准确率曲线', level=2)
    doc.add_paragraph('下图展示了CNN分支、Transformer分支以及集成模型在10轮训练中的准确率变化：')
    
    # 插入图片
    doc.add_picture('training_accuracy_curve.png', width=Inches(6.5))
    
    doc.add_heading('🔍 关键指标解读', level=2)
    
    # 关键指标表格
    metrics_table = doc.add_table(rows=1, cols=5)
    metrics_table.style = 'Table Grid'
    
    hdr_cells = metrics_table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '初始值 (Round 1)'
    hdr_cells[2].text = '峰值'
    hdr_cells[3].text = '稳定值 (Round 10)'
    hdr_cells[4].text = '趋势分析'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    metrics_data = [
        ('CNN分支', '71.1%', '93.8% (R7)', '91.3%', '快速收敛，局部特征提取稳定有效'),
        ('Transformer分支', '72.0%', '89.4% (R6)', '88.2%', '存在波动，全局建模需更多训练轮次稳定'),
        ('Ensemble集成', '71.0%', '90.4% (R7)', '89.1%', '融合策略有效平滑单模型波动，整体鲁棒性提升')
    ]
    
    for metric, initial, peak, stable, trend in metrics_data:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = initial
        row_cells[2].text = peak
        row_cells[3].text = stable
        row_cells[4].text = trend
    
    # 结果洞察
    doc.add_heading('💡 结果洞察', level=2)
    insights = [
        'CNN分支准确率显著高于Transformer分支，验证了局部形态特征对房颤检测的基础性作用',
        '集成模型在多数轮次优于单一Transformer，说明混合架构的互补性已初步体现',
        'Transformer分支在R3出现明显波动（80.1%），推测与可学习位置编码初期训练不稳定有关，后续可增加warmup策略'
    ]
    
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    
    # 四、配置参数备忘
    doc.add_heading('四、配置参数备忘', level=1)
    
    config = data["config"]
    
    p = doc.add_paragraph()
    p.add_run('训练配置参数：\n\n').bold = True
    for key, value in config.items():
        p.add_run(f'{key}: {value}\n')
    
    # 五、问题与下一步计划
    doc.add_heading('五、问题与下一步计划', level=1)
    
    doc.add_heading('⚠️ 当前问题', level=2)
    problems = [
        'Transformer分支训练波动较大，收敛速度慢于CNN分支',
        '集成策略目前为简单加权，未充分挖掘两分支的互补潜力'
    ]
    
    for problem in problems:
        doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_heading('🎯 下周计划', level=2)
    plans = [
        '引入学习率warmup + cosine decay策略，稳定Transformer训练初期',
        '尝试动态权重集成（如基于验证集表现的自适应融合）',
        '增加消融实验：对比 [CLS] vs GAP、可学习编码 vs 固定编码 的单独贡献',
        '在独立测试集上评估泛化能力，计算敏感性/特异性等临床指标'
    ]
    
    for plan in plans:
        doc.add_paragraph(plan, style='List Bullet')
    
    # 添加备注
    doc.add_paragraph()
    doc.add_paragraph('📌 备注：训练曲线图已通过代码生成并插入文档。如需调整图表样式或补充其他指标可视化，请随时告知。', style='Quote')
    
    doc.add_paragraph()
    doc.add_paragraph('祝工作顺利！🚀')
    
    # 保存文档
    doc.save('ConvTransformer_周报_2026年3月第2周.docx')
    print("Word文档已生成：ConvTransformer_周报_2026年3月第2周.docx")

# 执行生成
if __name__ == "__main__":
    generate_training_plot()
    generate_word_doc()
    print("\n✅ 所有文件生成完成！")